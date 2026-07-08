# core/parser.py — CV parsing logic for TalentIQ
# Strategy: try LLM first, fall back to hybrid deterministic parser

import re
import json
import logging
from typing import Dict, Any
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import config
from core.llm import get_llm_response

logger = logging.getLogger(__name__)


def extract_text_from_file(uploaded_file) -> str:
    """Extract raw text from PDF or DOCX uploaded file."""
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return _extract_pdf(uploaded_file)
    elif filename.endswith(".docx"):
        return _extract_docx(uploaded_file)
    else:
        return uploaded_file.read().decode("utf-8", errors="ignore")


def _extract_pdf(uploaded_file) -> str:
    try:
        import pdfplumber
        import io
        uploaded_file.seek(0)
        with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
            return "\n".join(
                page.extract_text() or "" for page in pdf.pages
            )
    except Exception as e:
        logger.error(f"Erro ao extrair PDF: {e}")
        return ""


def _extract_docx(uploaded_file) -> str:
    try:
        from docx import Document
        import io
        uploaded_file.seek(0)
        doc = Document(io.BytesIO(uploaded_file.read()))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"Erro ao extrair DOCX: {e}")
        return ""


def parse_cv(text: str) -> Dict[str, Any]:
    """
    Parse CV text and return structured candidate profile.
    Uses LLM if available, otherwise falls back to deterministic hybrid parser.
    """
    if config.LLM_ENGINE != "deterministic":
        result = _parse_with_llm(text)
        if result:
            result["metodo_extracao"] = f"LLM ({config.LLM_ENGINE})"
            return result

    result = _parse_deterministic(text)
    result["metodo_extracao"] = "Determinístico (Híbrido)"
    return result


def _parse_with_llm(text: str) -> Dict[str, Any]:
    system = (
        "Você é um especialista em recursos humanos. "
        "Extraia informações estruturadas de CVs. "
        "Responda APENAS com JSON válido, sem texto adicional."
    )
    prompt = f"""
Analise este CV e extraia as seguintes informações em JSON:
{{
  "nome": "nome completo",
  "email": "email ou null",
  "telefone": "telefone ou null",
  "competencias": ["lista de competências e habilidades"],
  "experiencia_anos": número inteiro de anos de experiência profissional mencionados EXPLICITAMENTE no CV (ex: "8 anos de experiência" → 8). NÃO calcule nem estime — se o CV não indicar um número explícito, some os períodos listados nos cargos. Retorne sempre um inteiro.
  "cargos_anteriores": ["lista de cargos ocupados"],
  "formacao": ["lista de formações académicas"],
  "idiomas": ["lista de idiomas"],
  "resumo": "resumo profissional em 2 frases"
}}

CV:
{text[:4000]}
"""
    try:
        response = get_llm_response(prompt, system)
        if response:
            clean = response.strip().strip("```json").strip("```").strip()
            return json.loads(clean)
    except Exception as e:
        logger.warning(f"LLM parsing falhou, usando fallback: {e}")
    return {}


def parse_tor(text: str) -> Dict[str, Any]:
    """
    Parse a Terms of Reference document and extract structured job vacancy data.
    Uses LLM if available, otherwise falls back to deterministic hybrid parser.
    """
    if config.LLM_ENGINE != "deterministic":
        result = _parse_tor_with_llm(text)
        if result:
            result["metodo_extracao"] = f"LLM ({config.LLM_ENGINE})"
            return result

    result = _parse_tor_deterministic(text)
    result["metodo_extracao"] = "Determinístico (Híbrido)"
    return result


def _parse_tor_with_llm(text: str) -> Dict[str, Any]:
    system = (
        "Você é um especialista em recursos humanos e recrutamento. "
        "Extraia informações estruturadas de Termos de Referência (ToR) de vagas de emprego. "
        "Responda APENAS com JSON válido, sem texto adicional."
    )
    prompt = f"""
Analise este Termo de Referência e extraia as seguintes informações em JSON:
{{
  "titulo": "título exacto da posição",
  "departamento": "departamento ou área funcional",
  "local": "localização da posição",
  "modalidade": "Presencial | Remoto | Híbrido",
  "nivel_formacao": "licenciatura | mestrado | doutoramento | curso técnico | outro",
  "anos_experiencia_min": número inteiro de anos mínimos de experiência,
  "tipo_contrato": "Tempo Inteiro | Tempo Parcial | Consultoria | Estágio",
  "salario": "faixa salarial mencionada ou null",
  "competencias_requeridas": ["lista de competências técnicas e transversais requeridas"],
  "responsabilidades": ["lista das principais responsabilidades"],
  "descricao": "resumo do contexto e objectivo da posição em 3-4 frases",
  "organizacao": "nome da organização que recruta",
  "prazo_candidatura": "data limite ou null"
}}

Termo de Referência:
{text[:5000]}
"""
    try:
        response = get_llm_response(prompt, system)
        if response:
            clean = response.strip().strip("```json").strip("```").strip()
            return json.loads(clean)
    except Exception as e:
        logger.warning(f"LLM ToR parsing falhou, usando fallback: {e}")
    return {}


def _parse_tor_deterministic(text: str) -> Dict[str, Any]:
    """Hybrid deterministic ToR parser."""
    text_lower = text.lower()
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    # Title — look for common ToR title patterns
    titulo = ""
    for i, line in enumerate(lines[:20]):
        if any(kw in line.lower() for kw in ["analista", "gestor", "gerente", "coordenador",
                                               "director", "especialista", "técnico", "oficial",
                                               "consultor", "assistente", "supervisor"]):
            titulo = line.strip()
            break
    if not titulo and lines:
        titulo = lines[1] if len(lines) > 1 else lines[0]

    # Organisation name — skip generic document headers, look for org identifiers
    _doc_headers = {"termos de referência", "terms of reference", "tor", "job description",
                    "descrição de cargo", "aviso de vaga", "anúncio de vaga"}
    organizacao = ""
    for line in lines[:10]:
        if line.lower().strip() not in _doc_headers and len(line) > 3 and len(line) < 80:
            organizacao = line
            break

    # Skills
    competencias = [kw for kw in config.SKILLS_KEYWORDS if kw.lower() in text_lower]

    # Education level
    nivel_formacao = ""
    for nivel in ["doutoramento", "phd", "mestrado", "mba", "licenciatura",
                   "bacharel", "curso técnico", "certificação"]:
        if nivel in text_lower:
            nivel_formacao = nivel
            break

    # Years of experience — regex: "X anos" or "X years"
    exp_match = re.search(r"(\d+)\s*(?:anos?|years?)\s*(?:de\s+)?(?:experi[eê]ncia)?", text_lower)
    anos_experiencia_min = int(exp_match.group(1)) if exp_match else 0

    # Location
    local = ""
    for loc in ["maputo", "beira", "nampula", "angola", "luanda", "brasil",
                 "são paulo", "lisboa", "portugal", "mozambique", "moçambique"]:
        if loc in text_lower:
            local = loc.title()
            break

    # Modality
    modalidade = "Presencial"
    if "remoto" in text_lower or "remote" in text_lower:
        modalidade = "Remoto"
    elif "híbrido" in text_lower or "hybrid" in text_lower:
        modalidade = "Híbrido"

    # Contract type
    tipo_contrato = "Tempo Inteiro"
    if "parcial" in text_lower or "part-time" in text_lower:
        tipo_contrato = "Tempo Parcial"
    elif "consultoria" in text_lower or "consulting" in text_lower:
        tipo_contrato = "Consultoria"
    elif "estágio" in text_lower or "internship" in text_lower:
        tipo_contrato = "Estágio"

    # Deadline
    prazo_match = re.search(r"\d{1,2}\s+(?:de\s+)?(?:janeiro|fevereiro|março|abril|maio|junho|"
                             r"julho|agosto|setembro|outubro|novembro|dezembro)\s+(?:de\s+)?\d{4}", text_lower)
    prazo = prazo_match.group(0).title() if prazo_match else None

    # Description — first substantial paragraph
    descricao = ""
    for line in lines:
        if len(line) > 100:
            descricao = line[:400]
            break

    # Department
    departamento = ""
    for dep in ["recursos humanos", "tecnologia", "finanças", "saúde pública",
                 "operações", "marketing", "vendas", "administração"]:
        if dep in text_lower:
            departamento = dep.title()
            break

    return {
        "titulo": titulo,
        "organizacao": organizacao,
        "departamento": departamento,
        "local": local,
        "modalidade": modalidade,
        "nivel_formacao": nivel_formacao,
        "anos_experiencia_min": anos_experiencia_min,
        "tipo_contrato": tipo_contrato,
        "salario": None,
        "competencias_requeridas": list(set(competencias)),
        "responsabilidades": [],
        "descricao": descricao,
        "prazo_candidatura": prazo,
    }


def _parse_deterministic(text: str) -> Dict[str, Any]:
    """Hybrid deterministic parser: keyword matching + regex rules."""
    text_lower = text.lower()

    # Extract skills via keyword matching
    competencias = [
        kw for kw in config.SKILLS_KEYWORDS
        if kw.lower() in text_lower
    ]

    # Extract education via keyword matching
    formacao = [
        kw.title() for kw in config.EDUCATION_KEYWORDS
        if kw.lower() in text_lower
    ]

    # Extract job titles via keyword matching
    cargos = [
        kw.title() for kw in config.JOB_TITLE_KEYWORDS
        if kw.lower() in text_lower
    ]

    # Extract email via regex
    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    email = email_match.group(0) if email_match else None

    # Extract phone via regex — must start with +, (, or 0, or be 9+ digits
    # Excludes year ranges like "2018-2024" by requiring non-digit boundary context
    phone_match = re.search(
        r"(?<!\d)(\+\d[\d\s\-\(\)]{7,15}\d|(?:\(?\d{2,3}\)?[\s\-])[\d\s\-]{6,12}\d)(?!\d)",
        text
    )
    telefone = phone_match.group(0).strip() if phone_match else None

    # Estimate years of experience
    # Strategy 1: explicit statement like "5 anos de experiência"
    import datetime as _dt
    _current_year = _dt.datetime.now().year
    explicit_exp = re.search(r"(\d+)\s*anos?\s*de\s*experi[eê]ncia", text_lower)
    if explicit_exp:
        experiencia_anos = int(explicit_exp.group(1))
    else:
        # Strategy 2: sum date ranges found in the text (e.g. "2021 - 2026", "2019–2021")
        # This avoids inflating experience by using the earliest education year
        ranges = re.findall(r"\b(20\d{2})\s*[-–—]\s*(20\d{2}|presente|atual|current|present)\b",
                            text_lower)
        total = 0
        for start_s, end_s in ranges:
            start = int(start_s)
            end = _current_year if end_s in ("presente", "atual", "current", "present") else int(end_s)
            if 2000 <= start <= _current_year and end >= start:
                total += end - start
        experiencia_anos = total if total > 0 else 0

    # Extract name — skip generic CV headers
    _header_noise = {"curriculum vitae", "cv", "resume", "candidatura", "perfil",
                     "dados pessoais", "personal information", "biodata"}
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    nome = "Desconhecido"
    for line in lines[:8]:
        if line.lower() not in _header_noise and len(line.split()) >= 2 and len(line) < 60:
            nome = line
            break

    # Extract languages
    idiomas = [
        kw.title() for kw in ["inglês", "português", "francês", "espanhol",
                               "english", "french", "spanish"]
        if kw in text_lower
    ]

    return {
        "nome": nome,
        "email": email,
        "telefone": telefone,
        "competencias": list(set(competencias)),
        "experiencia_anos": experiencia_anos,
        "cargos_anteriores": list(set(cargos)),
        "formacao": list(set(formacao)),
        "idiomas": list(set(idiomas)),
        "resumo": f"Perfil extraído automaticamente. {len(competencias)} competências identificadas.",
    }
