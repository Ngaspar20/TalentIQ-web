import os
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.contrib import messages
from django.views.decorators.http import require_POST
from .models import Candidato


def org_candidatos(request):
    return Candidato.objects.filter(organisation=request.user.organisation)


def candidato_list(request):
    candidatos = org_candidatos(request).select_related("vaga").order_by("-created_at")
    return render(request, "candidatos/list.html", {"candidatos": candidatos})


def candidato_create(request):
    from vagas.models import Vaga
    if request.method == "POST":
        vaga_id = request.POST.get("vaga_id", "").strip()
        nome = request.POST.get("nome", "").strip()
        if not nome:
            messages.error(request, "O nome do candidato Ã© obrigatÃ³rio.")
            vagas = Vaga.objects.filter(organisation=request.user.organisation, estado="Aberta")
            return render(request, "candidatos/create.html", {"vagas": vagas})

        import json
        competencias_raw = request.POST.get("competencias_input", "")
        competencias = [c.strip().lower() for c in competencias_raw.split(",") if c.strip()]

        formacao_raw = request.POST.get("formacao_input", "")
        formacao = [f.strip() for f in formacao_raw.splitlines() if f.strip()]

        idiomas_raw = request.POST.get("idiomas_input", "")
        idiomas = [i.strip() for i in idiomas_raw.split(",") if i.strip()]

        vaga = None
        if vaga_id:
            try:
                vaga = Vaga.objects.get(pk=vaga_id, organisation=request.user.organisation)
            except Vaga.DoesNotExist:
                pass

        candidato = Candidato.objects.create(
            organisation=request.user.organisation,
            vaga=vaga,
            nome=nome,
            email=request.POST.get("email", "").strip(),
            telefone=request.POST.get("telefone", "").strip(),
            experiencia_anos=int(request.POST.get("experiencia_anos", 0) or 0),
            competencias=competencias,
            formacao=formacao,
            idiomas=idiomas,
            resumo=request.POST.get("resumo", "").strip(),
            notas=request.POST.get("notas", "").strip(),
            created_by=request.user,
        )
        messages.success(request, f"Candidato '{candidato.nome}' adicionado com sucesso!")
        return redirect("candidato_list")

    vagas = Vaga.objects.filter(organisation=request.user.organisation, estado="Aberta")
    return render(request, "candidatos/create.html", {"vagas": vagas})


def candidato_detail(request, pk):
    candidato = get_object_or_404(org_candidatos(request), pk=pk)
    from .models import NotaEntrevista, AvaliacaoSession
    nota = NotaEntrevista.objects.filter(candidato=candidato).first()
    avaliacao_session = candidato.avaliacao_sessions.order_by("-created_at").first()
    return render(request, "candidatos/detail.html", {
        "candidato": candidato,
        "nota": nota,
        "avaliacao_session": avaliacao_session,
    })


def candidato_edit(request, pk):
    from vagas.models import Vaga
    candidato = get_object_or_404(org_candidatos(request), pk=pk)
    vagas = Vaga.objects.filter(organisation=request.user.organisation).order_by("titulo")
    if request.method == "POST":
        nome = request.POST.get("nome", "").strip()
        if not nome:
            messages.error(request, "O nome do candidato Ã© obrigatÃ³rio.")
            return render(request, "candidatos/edit.html", {"candidato": candidato, "vagas": vagas})

        competencias_raw = request.POST.get("competencias_input", "")
        competencias = [c.strip().lower() for c in competencias_raw.split(",") if c.strip()]

        formacao_raw = request.POST.get("formacao_input", "")
        formacao = [f.strip() for f in formacao_raw.splitlines() if f.strip()]

        idiomas_raw = request.POST.get("idiomas_input", "")
        idiomas = [i.strip() for i in idiomas_raw.split(",") if i.strip()]

        vaga_id = request.POST.get("vaga_id", "").strip()
        if vaga_id:
            try:
                candidato.vaga = Vaga.objects.get(pk=vaga_id, organisation=request.user.organisation)
            except Vaga.DoesNotExist:
                candidato.vaga = None
        else:
            candidato.vaga = None

        candidato.nome = nome
        candidato.email = request.POST.get("email", "").strip()
        candidato.telefone = request.POST.get("telefone", "").strip()
        candidato.experiencia_anos = int(request.POST.get("experiencia_anos", 0) or 0)
        candidato.etapa = request.POST.get("etapa", "Candidatura Recebida")
        candidato.competencias = competencias
        candidato.formacao = formacao
        candidato.idiomas = idiomas
        candidato.resumo = request.POST.get("resumo", "").strip()
        candidato.notas = request.POST.get("notas", "").strip()
        candidato.save()
        messages.success(request, f"Candidato '{candidato.nome}' actualizado.")
        return redirect("candidato_detail", pk=candidato.pk)

    return render(request, "candidatos/edit.html", {"candidato": candidato, "vagas": vagas})


def candidato_delete(request, pk):
    candidato = get_object_or_404(org_candidatos(request), pk=pk)
    if request.method == "POST":
        candidato.delete()
        messages.success(request, f"Candidato '{candidato.nome}' eliminado.")
        return redirect("candidato_list")
    return render(request, "candidatos/confirm_delete.html", {"candidato": candidato})


CARTA_TIPOS = {
    "triagem_exclusao": "Carta de Exclusão de Triagem",
    "pre_selecao": "Carta de Pré-Selecção",
    "pos_entrevista_rejeicao": "Carta de Rejeição Pós-Entrevista",
    "oferta": "Carta de Oferta",
}


def _gerar_texto_carta(tipo, nome, vaga, org, data_entrevista=""):
    from core.llm import get_llm_response

    info_entrevista = f" A entrevista está agendada para {data_entrevista}." if data_entrevista else " A nossa equipa de RH entrará em contacto brevemente com os detalhes da entrevista."

    prompts = {
        "triagem_exclusao": f"Escreve uma carta profissional e respeitosa de exclusão de triagem em português europeu para {nome}, que se candidatou ao cargo de {vaga} na {org}. O candidato não foi seleccionado para a fase de entrevista. Sê respeitoso, agradecido e conciso (3 parágrafos). Escreve apenas o corpo da carta a começar com 'Exmo./Exma. Sr./Sra. {nome},'.",
        "pre_selecao": f"Escreve uma carta profissional de pré-selecção em português europeu para {nome}, convidando-o/a para entrevista para o cargo de {vaga} na {org}.{info_entrevista} Felicita o candidato pela selecção. Sê caloroso e conciso (2-3 parágrafos). Escreve apenas o corpo da carta a começar com 'Exmo./Exma. Sr./Sra. {nome},'.",
        "pos_entrevista_rejeicao": f"Escreve uma carta profissional de rejeição pós-entrevista em português europeu para {nome}, que foi entrevistado/a para o cargo de {vaga} na {org}. Agradece a participação, reconhece o esforço, informa que outro candidato foi seleccionado. Sê caloroso e respeitoso (3 parágrafos). Escreve apenas o corpo da carta a começar com 'Exmo./Exma. Sr./Sra. {nome},'.",
        "oferta": f"Escreve uma carta profissional de oferta de emprego em português europeu para {nome} para o cargo de {vaga} na {org}. Felicita o candidato pela selecção, menciona que os detalhes completos serão enviados em breve. Sê caloroso e formal (2-3 parágrafos). Escreve apenas o corpo da carta a começar com 'Exmo./Exma. Sr./Sra. {nome},'.",
    }
    system = "És um profissional de recursos humanos a escrever cartas formais de emprego em português europeu. Sê profissional, caloroso e conciso."
    texto = get_llm_response(prompts[tipo], system)
    if texto:
        return texto

    info_fb = f"A entrevista está agendada para {data_entrevista}." if data_entrevista else "A nossa equipa de RH entrará em contacto brevemente para confirmar os detalhes da entrevista."

    fallbacks = {
        "triagem_exclusao": f"Exmo./Exma. Sr./Sra. {nome},\n\nVimos por este meio agradecer a sua candidatura ao cargo de {vaga} na {org}. Agradecemos o tempo e esforço investidos no processo de candidatura.\n\nApós análise cuidada de todas as candidaturas recebidas, informamos com pesar que não foi possível seleccioná-lo/a para a fase de entrevista. Esta foi uma decisão difícil, dado o elevado número de candidatos qualificados que se candidataram.\n\nDesejamos-lhe muito sucesso na sua procura de emprego e esperamos que considere candidatar-se a futuras oportunidades na nossa organização.\n\nCom os melhores cumprimentos,\n\nRecursos Humanos\n{org}",
        "pre_selecao": f"Exmo./Exma. Sr./Sra. {nome},\n\nTem o prazer de informar que, após análise das candidaturas recebidas para o cargo de {vaga} na {org}, foi seleccionado/a para prosseguir para a fase de entrevista.\n\n{info_fb}\n\nAgradecemos o seu interesse em fazer parte da {org} e aguardamos com expectativa conhecê-lo/a pessoalmente.\n\nCom os melhores cumprimentos,\n\nRecursos Humanos\n{org}",
        "pos_entrevista_rejeicao": f"Exmo./Exma. Sr./Sra. {nome},\n\nVimos por este meio agradecer a sua participação na entrevista para o cargo de {vaga} na {org}. Agradecemos o tempo e esforço que dedicou ao processo de selecção.\n\nApós deliberação cuidada, decidimos avançar com outro candidato cujo perfil correspondeu mais especificamente aos requisitos do cargo. Esta foi uma decisão difícil, dado o elevado nível dos candidatos entrevistados.\n\nEncorajamo-lo/a a candidatar-se a futuras oportunidades na {org} e desejamos-lhe muito sucesso na sua carreira.\n\nCom os melhores cumprimentos,\n\nRecursos Humanos\n{org}",
        "oferta": f"Exmo./Exma. Sr./Sra. {nome},\n\nEm nome da {org}, é com grande satisfação que lhe comunicamos a sua selecção para o cargo de {vaga}. Após um processo de selecção criterioso, o júri de selecção concluiu unanimemente que o seu perfil e experiência fazem de si o/a candidato/a ideal para esta função.\n\nA nossa equipa de Recursos Humanos entrará em contacto brevemente com todos os detalhes da proposta, incluindo termos, condições e data de início.\n\nAguardamos com entusiasmo a sua integração na nossa equipa.\n\nCom os melhores cumprimentos,\n\nRecursos Humanos\n{org}",
    }
    return fallbacks[tipo]


def gerar_carta(request, pk):
    candidato = get_object_or_404(org_candidatos(request), pk=pk)
    tipo = request.GET.get("tipo", "")
    if tipo not in CARTA_TIPOS:
        return redirect("candidato_detail", pk=pk)

    # Pre-selection letter needs interview date — show form first on GET
    if tipo == "pre_selecao" and request.method == "GET" and not request.GET.get("gerar"):
        return render(request, "candidatos/carta_data_entrevista.html", {
            "candidato": candidato,
            "tipo": tipo,
            "tipo_label": CARTA_TIPOS[tipo],
        })

    nome = candidato.nome
    vaga = candidato.vaga.titulo if candidato.vaga else "o cargo anunciado"
    org = candidato.organisation.name if hasattr(candidato, 'organisation') and candidato.organisation else "a nossa organização"
    data_raw = request.POST.get("data_entrevista", "") or request.GET.get("data_entrevista", "")
    hora = request.POST.get("hora_entrevista", "")
    formato = request.POST.get("formato_entrevista", "presencial")
    local = request.POST.get("local_entrevista", "")
    plataforma = request.POST.get("plataforma_virtual", "")
    link = request.POST.get("link_virtual", "")
    if formato == "virtual":
        local_info = f"entrevista virtual via {plataforma}"
        if link:
            local_info += f" ({link})"
    else:
        local_info = local
    parts = [p for p in [data_raw, hora, local_info] if p]
    data_entrevista = ", ".join(parts)

    texto = _gerar_texto_carta(tipo, nome, vaga, org, data_entrevista)
    request.session[f"carta_{tipo}_{pk}"] = texto

    return render(request, "candidatos/carta_preview.html", {
        "candidato": candidato,
        "tipo": tipo,
        "tipo_label": CARTA_TIPOS[tipo],
        "texto": texto,
    })


def download_carta(request, pk):
    candidato = get_object_or_404(org_candidatos(request), pk=pk)
    tipo = request.GET.get("tipo", "")
    texto = request.POST.get("texto", "") or request.session.get(f"carta_{tipo}_{pk}", "")

    if not texto or tipo not in CARTA_TIPOS:
        return redirect("candidato_detail", pk=pk)

    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    from datetime import date

    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    org_name = candidato.organisation.name if candidato.organisation else "Organisation"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(org_name)
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(date.today().strftime("%d %B %Y")).font.size = Pt(11)

    doc.add_paragraph()

    vaga_titulo = candidato.vaga.titulo if candidato.vaga else "Position"
    p = doc.add_paragraph()
    run = p.add_run(f"Re: {vaga_titulo}")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    for line in texto.split('\n'):
        p = doc.add_paragraph(line if line.strip() else "")
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = candidato.nome.replace(' ', '_')
    filename = f"{tipo}_{safe_name}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def enviar_avaliacao_juri(request, pk):
    candidato = get_object_or_404(org_candidatos(request), pk=pk)
    if request.method != "POST":
        return redirect("candidato_detail", pk=pk)
    from .models import AvaliacaoSession
    chair_email = request.POST.get("chair_email", "").strip()
    session = AvaliacaoSession.objects.create(candidato=candidato, chair_email=chair_email)
    link = request.build_absolute_uri(f"/candidatos/avaliacao/{session.token}/")
    return render(request, "candidatos/avaliacao_link.html", {
        "candidato": candidato,
        "session": session,
        "link": link,
    })


def avaliacao_juri_view(request, token):
    from .models import AvaliacaoSession
    session = get_object_or_404(AvaliacaoSession, token=token)
    if session.estado == AvaliacaoSession.ESTADO_CONFIRMADA:
        return render(request, "candidatos/avaliacao_fechada.html", {"session": session})
    if request.method == "POST":
        session.data_entrevista = request.POST.get("data_entrevista") or None
        p = request.POST.get("pontuacao", "")
        session.pontuacao = int(p) if p.isdigit() and 1 <= int(p) <= 5 else None
        session.recomendacao = request.POST.get("recomendacao", "")
        session.pontos_fortes = request.POST.get("pontos_fortes", "").strip()
        session.pontos_fracos = request.POST.get("pontos_fracos", "").strip()
        session.notas = request.POST.get("notas", "").strip()
        session.estado = AvaliacaoSession.ESTADO_SUBMETIDA
        session.save()
        return render(request, "candidatos/avaliacao_obrigado.html", {"session": session})
    return render(request, "candidatos/avaliacao_juri.html", {
        "session": session,
        "candidato": session.candidato,
    })


def avaliacao_confirmar(request, pk, session_id):
    candidato = get_object_or_404(org_candidatos(request), pk=pk)
    from .models import AvaliacaoSession
    session = get_object_or_404(AvaliacaoSession, pk=session_id, candidato=candidato)
    if request.method != "POST":
        return redirect("candidato_detail", pk=pk)

    session.estado = AvaliacaoSession.ESTADO_CONFIRMADA
    session.save()

    # Copy evaluation to NotaEntrevista for persistent display
    from .models import NotaEntrevista
    nota, _ = NotaEntrevista.objects.get_or_create(candidato=candidato)
    nota.data_entrevista = session.data_entrevista
    nota.pontuacao = session.pontuacao
    nota.recomendacao = session.recomendacao
    nota.pontos_fortes = session.pontos_fortes
    nota.pontos_fracos = session.pontos_fracos
    nota.notas = session.notas
    nota.save()

    # Auto-transition candidate stage
    rec = session.recomendacao
    if rec == "recomendado":
        candidato.etapa = "Proposta"
        messages.success(request, f"{candidato.nome} avançou para Proposta. Carta de oferta disponível.")
    elif rec == "nao_recomendado":
        candidato.etapa = "Rejeitado"
        messages.warning(request, f"{candidato.nome} marcado como Rejeitado. Carta de rejeição disponível.")
    else:
        messages.info(request, f"Avaliação confirmada. {candidato.nome} mantém-se em avaliação.")
    candidato.save()

    return redirect("candidato_detail", pk=pk)


@require_POST
def guardar_nota_entrevista(request, pk):
    candidato = get_object_or_404(org_candidatos(request), pk=pk)
    from .models import NotaEntrevista

    nota, _ = NotaEntrevista.objects.get_or_create(candidato=candidato)
    nota.data_entrevista = request.POST.get("data_entrevista") or None
    pontuacao = request.POST.get("pontuacao", "")
    nota.pontuacao = int(pontuacao) if pontuacao.isdigit() and 1 <= int(pontuacao) <= 5 else None
    nota.recomendacao = request.POST.get("recomendacao", "")
    nota.notas = request.POST.get("notas", "").strip()
    nota.pontos_fortes = request.POST.get("pontos_fortes", "").strip()
    nota.pontos_fracos = request.POST.get("pontos_fracos", "").strip()
    nota.save()
    messages.success(request, "Notas de entrevista guardadas.")
    return redirect("candidato_detail", pk=pk)


@require_POST
def parse_cv_view(request):
    """HTMX â€” extracts text from uploaded CV file and returns preview."""
    uploaded = request.FILES.get("cv_file")
    if not uploaded:
        return HttpResponse('<div class="alert-error">Nenhum ficheiro recebido.</div>')

    try:
        from core.parser import extract_text_from_file
        texto = extract_text_from_file(uploaded)
    except Exception as e:
        return HttpResponse(f'<div class="alert-error">Erro ao extrair texto: {e}</div>')

    if not texto.strip():
        return HttpResponse('<div class="alert-error">NÃ£o foi possÃ­vel extrair texto. O ficheiro pode ser uma imagem digitalizada.</div>')

    return render(request, "candidatos/_cv_preview.html", {
        "texto": texto[:4000],
        "texto_completo": texto,
    })


@require_POST
def analyse_cv_view(request):
    """HTMX â€” sends CV text to Grok and returns pre-filled candidate form."""
    from django.conf import settings
    texto = request.POST.get("texto_completo", "")
    if not texto.strip():
        return HttpResponse('<div class="alert-error">Texto nÃ£o encontrado. Carregue o CV novamente.</div>')

    try:
        os.environ["GROK_API_KEY"] = settings.GROK_API_KEY
        os.environ["LLM_ENGINE"] = settings.LLM_ENGINE

        from core.parser import parse_cv
        extraido = parse_cv(texto)
    except Exception as e:
        return HttpResponse(f'<div class="alert-error">Erro na anÃ¡lise IA: {e}</div>')

    return render(request, "candidatos/_cv_form_fields.html", {
        "cv": extraido,
        "metodo": extraido.get("metodo_extracao", "IA"),
    })

