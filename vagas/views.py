import json
import sys
import os
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, HttpResponse
from django.contrib import messages
from django.views.decorators.http import require_POST
from django.contrib.auth.decorators import login_required
from .models import Vaga

# Make sure core/ is on the path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


def org_vagas(request):
    return Vaga.objects.filter(organisation=request.user.organisation)


def vaga_list(request):
    vagas = org_vagas(request).order_by("-created_at")
    return render(request, "vagas/list.html", {"vagas": vagas})


def vaga_create(request):
    if request.method == "POST":
        titulo = request.POST.get("titulo", "").strip()
        if not titulo:
            messages.error(request, "O tÃ­tulo da vaga Ã© obrigatÃ³rio.")
            return render(request, "vagas/create.html")

        competencias_raw = request.POST.get("competencias_input", "")
        competencias = [c.strip().lower() for c in competencias_raw.split(",") if c.strip()]

        responsabilidades_raw = request.POST.get("responsabilidades_input", "")
        responsabilidades = [r.strip().lstrip("â€¢").strip() for r in responsabilidades_raw.splitlines() if r.strip()]

        tor_filename = request.POST.get("tor_filename", "").strip()

        vaga = Vaga.objects.create(
            organisation=request.user.organisation,
            titulo=titulo,
            organizacao=request.POST.get("organizacao", "").strip(),
            departamento=request.POST.get("departamento", "Outro"),
            local=request.POST.get("local", "").strip(),
            modalidade=request.POST.get("modalidade", "Presencial"),
            nivel_formacao=request.POST.get("nivel_formacao", "").lower(),
            anos_experiencia_min=int(request.POST.get("anos_experiencia_min", 0) or 0),
            tipo_contrato=request.POST.get("tipo_contrato", "Tempo Inteiro"),
            salario=request.POST.get("salario", "").strip(),
            prazo_candidatura=request.POST.get("prazo_candidatura", "").strip(),
            competencias_requeridas=competencias,
            responsabilidades=responsabilidades,
            descricao=request.POST.get("descricao", "").strip(),
            tor_file_path=tor_filename,
            tor_aprovado=False,
            origem="ToR" if tor_filename else "Manual",
            created_by=request.user,
        )
        messages.success(request, f"Vaga '{vaga.titulo}' criada com sucesso!")
        return redirect("vaga_list")

    return render(request, "vagas/create.html")


def vaga_detail(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    from candidatos.models import Candidato, AvaliacaoSession
    guiao_session = vaga.guiao_sessions.order_by("-created_at").first()

    todos = list(
        Candidato.objects
        .filter(vaga=vaga)
        .select_related("nota_entrevista")
        .order_by("-score_fit", "nome")
    )

    cands_triagem   = [c for c in todos if c.etapa in ("Candidatura Recebida", "Em Triagem")]
    cands_entrevista = [c for c in todos if c.etapa == "Entrevista"]
    cands_decisao   = [c for c in todos if c.etapa in ("Proposta", "Contratado", "Rejeitado")]

    # Build entrevista list with attached avaliacao_session (avoids underscore attr in templates)
    eval_session_map = {}
    if cands_entrevista:
        ids = [c.pk for c in cands_entrevista]
        for s in AvaliacaoSession.objects.filter(candidato_id__in=ids).order_by("candidato_id", "-created_at"):
            if s.candidato_id not in eval_session_map:
                eval_session_map[s.candidato_id] = s

    cands_entrevista_data = [
        {"candidato": c, "avaliacao_session": eval_session_map.get(c.pk),
         "nota": getattr(c, "nota_entrevista", None)}
        for c in cands_entrevista
    ]

    n_scored = sum(1 for c in todos if c.score_fit is not None)

    from .models import ComiteSession
    comite_sessions = list(vaga.comite_sessions.all())

    return render(request, "vagas/detail.html", {
        "vaga": vaga,
        "guiao_session": guiao_session,
        "cands_triagem": cands_triagem,
        "cands_entrevista_data": cands_entrevista_data,
        "cands_decisao": cands_decisao,
        "total_candidatos": len(todos),
        "n_entrevista": len(cands_entrevista),
        "n_scored": n_scored,
        "comite_sessions": comite_sessions,
    })


@require_POST
def vaga_confirmar_analise(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    vaga.tor_analisado = True
    vaga.save(update_fields=["tor_analisado"])
    messages.success(request, "Análise IA confirmada. Pode agora aprovar os Termos de Referência.")
    return redirect("vaga_detail", pk=pk)


@require_POST
def vaga_aprovar_tor(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    vaga.tor_aprovado = True
    vaga.save(update_fields=["tor_aprovado"])
    messages.success(request, f"Termos de Referência de '{vaga.titulo}' aprovados.")
    return redirect("vaga_detail", pk=pk)


def vaga_edit(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    if request.method == "POST":
        titulo = request.POST.get("titulo", "").strip()
        if not titulo:
            messages.error(request, "O tÃ­tulo da vaga Ã© obrigatÃ³rio.")
            return render(request, "vagas/edit.html", {"vaga": vaga})

        competencias_raw = request.POST.get("competencias_input", "")
        competencias = [c.strip().lower() for c in competencias_raw.split(",") if c.strip()]

        responsabilidades_raw = request.POST.get("responsabilidades_input", "")
        responsabilidades = [r.strip().lstrip("â€¢").strip() for r in responsabilidades_raw.splitlines() if r.strip()]

        vaga.titulo = titulo
        vaga.organizacao = request.POST.get("organizacao", "").strip()
        vaga.departamento = request.POST.get("departamento", "Outro")
        vaga.local = request.POST.get("local", "").strip()
        vaga.modalidade = request.POST.get("modalidade", "Presencial")
        vaga.estado = request.POST.get("estado", "Aberta")
        vaga.nivel_formacao = request.POST.get("nivel_formacao", "").lower()
        vaga.anos_experiencia_min = int(request.POST.get("anos_experiencia_min", 0) or 0)
        vaga.tipo_contrato = request.POST.get("tipo_contrato", "Tempo Inteiro")
        vaga.salario = request.POST.get("salario", "").strip()
        vaga.prazo_candidatura = request.POST.get("prazo_candidatura", "").strip()
        vaga.competencias_requeridas = competencias
        vaga.responsabilidades = responsabilidades
        vaga.descricao = request.POST.get("descricao", "").strip()
        vaga.save()
        messages.success(request, f"Vaga '{vaga.titulo}' actualizada.")
        return redirect("vaga_detail", pk=vaga.pk)

    return render(request, "vagas/edit.html", {"vaga": vaga})


def vaga_delete(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    if request.method == "POST":
        vaga.delete()
        messages.success(request, f"Vaga '{vaga.titulo}' eliminada.")
        return redirect("vaga_list")
    return render(request, "vagas/confirm_delete.html", {"vaga": vaga})


@require_POST
def parse_tor_view(request):
    """
    HTMX endpoint â€” called when user uploads a ToR file.
    Step 1: extract text and return preview.
    """
    from django.conf import settings
    uploaded = request.FILES.get("tor_file")
    if not uploaded:
        return HttpResponse('<div class="alert-error">Nenhum ficheiro recebido.</div>')

    try:
        from core.parser import extract_text_from_file
        texto = extract_text_from_file(uploaded)
    except Exception as e:
        return HttpResponse(f'<div class="alert-error">Erro ao extrair texto: {e}</div>')

    if not texto.strip():
        return HttpResponse('<div class="alert-error">NÃ£o foi possÃ­vel extrair texto. O ficheiro pode ser uma imagem digitalizada.</div>')

    # Return the text preview + hidden field + AI analysis button
    texto_preview = texto[:4000]
    return render(request, "vagas/_tor_preview.html", {
        "texto": texto_preview,
        "texto_completo": texto,
    })


@require_POST
def analyse_tor_view(request):
    """
    HTMX endpoint â€” called when user clicks 'Analisar com IA'.
    Sends text to Grok and returns pre-filled form fields.
    """
    from django.conf import settings
    texto = request.POST.get("texto_completo", "")
    if not texto.strip():
        return HttpResponse('<div class="alert-error">Texto nÃ£o encontrado. Carregue o ficheiro novamente.</div>')

    try:
        os.environ["GROK_API_KEY"] = settings.GROK_API_KEY
        os.environ["LLM_ENGINE"] = settings.LLM_ENGINE

        from core.parser import parse_tor
        extraido = parse_tor(texto)
    except Exception as e:
        return HttpResponse(f'<div class="alert-error">Erro na anÃ¡lise IA: {e}</div>')

    return render(request, "vagas/_form_fields.html", {
        "tor": extraido,
        "metodo": extraido.get("metodo_extracao", "IA"),
    })


def gerar_perguntas_entrevista(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)

    competencias = ", ".join(vaga.competencias_requeridas) if vaga.competencias_requeridas else "nao especificadas"
    responsabilidades = "; ".join(vaga.responsabilidades) if vaga.responsabilidades else "nao especificadas"
    descricao = vaga.descricao or ""

    from core.llm import get_llm_response

    prompt = f"""Gera um guiao estruturado de perguntas de entrevista em portugues europeu para o cargo de {vaga.titulo}.

Informacao da vaga:
- Cargo: {vaga.titulo}
- Departamento: {vaga.departamento}
- Formacao minima: {vaga.nivel_formacao or 'nao especificada'}
- Experiencia minima: {vaga.anos_experiencia_min} anos
- Competencias requeridas: {competencias}
- Responsabilidades: {responsabilidades}
- Descricao: {descricao[:500] if descricao else 'nao disponivel'}

Usa EXACTAMENTE este formato para cada categoria e pergunta (sem adicionar mais texto):

TOTAL DE PERGUNTAS: exactamente 10, distribuidas assim:

## APRESENTACAO E MOTIVACAO
P: [pergunta completa]
A: [o que avaliar na resposta, em 1 frase]
P: [pergunta completa]
A: [o que avaliar na resposta, em 1 frase]

## EXPERIENCIA E HISTORIAL PROFISSIONAL
P: [pergunta completa]
A: [o que avaliar]
P: [pergunta completa]
A: [o que avaliar]

## COMPETENCIAS TECNICAS
P: [pergunta tecnica especifica ao cargo de {vaga.titulo}]
A: [o que avaliar]
P: [pergunta tecnica especifica ao cargo de {vaga.titulo}]
A: [o que avaliar]
P: [pergunta tecnica especifica ao cargo de {vaga.titulo}]
A: [o que avaliar]

## COMPETENCIAS COMPORTAMENTAIS
P: [pergunta comportamental]
A: [o que avaliar]
P: [pergunta comportamental]
A: [o que avaliar]

## SITUACOES HIPOTETICAS
P: [situacao hipotetica relevante para o cargo]
A: [o que avaliar]

## QUESTOES DO CANDIDATO
P: Dar espaco ao/a candidato/a para colocar questoes sobre o cargo e a organizacao.
A: Interesse genuino, qualidade e pertinencia das questoes colocadas."""

    system = "Es um especialista em recursos humanos e seleccao de pessoal. Escreve em portugues europeu formal. Segue o formato pedido rigorosamente."
    texto = get_llm_response(prompt, system)

    if not texto:
        if vaga.competencias_requeridas:
            c1 = vaga.competencias_requeridas[0]
            c2 = vaga.competencias_requeridas[1] if len(vaga.competencias_requeridas) > 1 else c1
            comp_perguntas = (
                f"P: Descreva a sua experiencia com {c1} e como a aplicou em contexto profissional.\n"
                f"A: Profundidade de conhecimento, exemplos concretos, relevancia para o cargo.\n"
                f"P: Como utilizaria {c2} nas responsabilidades diarias deste cargo?\n"
                f"A: Aplicacao pratica, raciocinio tecnico, alinhamento com a funcao.\n"
                f"P: Como se mantém actualizado/a nas tendencias da sua area profissional?\n"
                f"A: Curiosidade intelectual, iniciativa de aprendizagem continua.\n"
            )
        else:
            comp_perguntas = (
                "P: Quais sao as suas principais competencias tecnicas relevantes para este cargo?\n"
                "A: Alinhamento com os requisitos, profundidade de conhecimento.\n"
                "P: Descreva uma situacao em que teve de aprender rapidamente uma nova ferramenta ou metodologia.\n"
                "A: Capacidade de aprendizagem, adaptacao, proactividade.\n"
                "P: Como garante a qualidade do seu trabalho tecnico?\n"
                "A: Metodo, atencao ao detalhe, orientacao para resultados.\n"
            )

        texto = f"""## APRESENTACAO E MOTIVACAO
P: Apresente-se brevemente e descreva o seu percurso profissional.
A: Capacidade de sintese, clareza de comunicacao, coerencia do percurso.
P: O que o/a motivou a candidatar-se a este cargo na nossa organizacao?
A: Conhecimento da organizacao, motivacao genuina, alinhamento de valores.

## EXPERIENCIA E HISTORIAL PROFISSIONAL
P: Descreva a sua experiencia mais relevante para o cargo de {vaga.titulo}.
A: Alinhamento com os requisitos da vaga, profundidade e qualidade da experiencia.
P: Qual foi o maior desafio profissional que enfrentou e como o resolveu?
A: Capacidade de resolucao de problemas, resiliencia, aprendizagem com a experiencia.

## COMPETENCIAS TECNICAS
{comp_perguntas}
## COMPETENCIAS COMPORTAMENTAIS
P: Como gere situacoes de conflito com colegas ou superiores hierarquicos?
A: Inteligencia emocional, comunicacao assertiva, capacidade de mediar.
P: Descreva uma situacao em que teve de trabalhar sob pressao e com prazos apertados.
A: Resistencia ao stress, organizacao, capacidade de priorizar.

## SITUACOES HIPOTETICAS
P: Se tivesse de gerir varias tarefas urgentes em simultaneo, como procederia?
A: Gestao de prioridades, metodologia de trabalho, pedido de apoio quando necessario.

## QUESTOES DO CANDIDATO
P: Dar espaco ao/a candidato/a para colocar questoes sobre o cargo e a organizacao.
A: Interesse genuino, qualidade e pertinencia das questoes colocadas."""

    request.session[f"perguntas_{pk}"] = texto

    categorias_parsed = _parse_perguntas(texto)
    categorias = [
        {"nome": nome, "perguntas": rows}
        for nome, rows in categorias_parsed
    ]

    return render(request, "vagas/perguntas_preview.html", {
        "vaga": vaga,
        "texto": texto,
        "categorias": categorias,
    })


def enviar_guiao_juri(request, pk):
    """HR generates a tokenized link and sends to committee chair."""
    vaga = get_object_or_404(org_vagas(request), pk=pk)

    if request.method != "POST":
        return redirect("vaga_detail", pk=pk)

    from core.llm import get_llm_response

    competencias = ", ".join(vaga.competencias_requeridas) if vaga.competencias_requeridas else "nao especificadas"
    responsabilidades = "; ".join(vaga.responsabilidades) if vaga.responsabilidades else "nao especificadas"
    descricao = vaga.descricao or ""

    prompt = f"""Gera um guiao estruturado de perguntas de entrevista em portugues europeu para o cargo de {vaga.titulo}.

Informacao da vaga:
- Cargo: {vaga.titulo}
- Departamento: {vaga.departamento}
- Formacao minima: {vaga.nivel_formacao or 'nao especificada'}
- Experiencia minima: {vaga.anos_experiencia_min} anos
- Competencias requeridas: {competencias}
- Responsabilidades: {responsabilidades}
- Descricao: {descricao[:500] if descricao else 'nao disponivel'}

TOTAL DE PERGUNTAS: exactamente 10, distribuidas assim:

## APRESENTACAO E MOTIVACAO
P: [pergunta completa]
A: [o que avaliar na resposta, em 1 frase]
P: [pergunta completa]
A: [o que avaliar na resposta, em 1 frase]

## EXPERIENCIA E HISTORIAL PROFISSIONAL
P: [pergunta completa]
A: [o que avaliar]
P: [pergunta completa]
A: [o que avaliar]

## COMPETENCIAS TECNICAS
P: [pergunta tecnica especifica ao cargo de {vaga.titulo}]
A: [o que avaliar]
P: [pergunta tecnica especifica ao cargo de {vaga.titulo}]
A: [o que avaliar]
P: [pergunta tecnica especifica ao cargo de {vaga.titulo}]
A: [o que avaliar]

## COMPETENCIAS COMPORTAMENTAIS
P: [pergunta comportamental]
A: [o que avaliar]
P: [pergunta comportamental]
A: [o que avaliar]

## SITUACOES HIPOTETICAS
P: [situacao hipotetica relevante para o cargo]
A: [o que avaliar]

## QUESTOES DO CANDIDATO
P: Dar espaco ao/a candidato/a para colocar questoes sobre o cargo e a organizacao.
A: Interesse genuino, qualidade e pertinencia das questoes colocadas."""

    system = "Es um especialista em recursos humanos e seleccao de pessoal. Escreve em portugues europeu formal. Segue o formato pedido rigorosamente."
    texto = get_llm_response(prompt, system) or _fallback_texto(vaga)

    chair_email = request.POST.get("chair_email", "").strip()

    from .models import InterviewGuideSession
    session = InterviewGuideSession.objects.create(
        vaga=vaga,
        texto_gerado=texto,
        chair_email=chair_email,
    )

    link = request.build_absolute_uri(f"/vagas/juri/{session.token}/")
    messages.success(request, f"Link gerado com sucesso.")
    return render(request, "vagas/guiao_link.html", {
        "vaga": vaga,
        "session": session,
        "link": link,
    })


def _fallback_texto(vaga):
    if vaga.competencias_requeridas:
        c1 = vaga.competencias_requeridas[0]
        c2 = vaga.competencias_requeridas[1] if len(vaga.competencias_requeridas) > 1 else c1
        comp = (
            f"P: Descreva a sua experiencia com {c1} e como a aplicou em contexto profissional.\n"
            f"A: Profundidade de conhecimento, exemplos concretos, relevancia para o cargo.\n"
            f"P: Como utilizaria {c2} nas responsabilidades diarias deste cargo?\n"
            f"A: Aplicacao pratica, raciocinio tecnico, alinhamento com a funcao.\n"
            f"P: Como se mantém actualizado/a nas tendencias da sua area profissional?\n"
            f"A: Curiosidade intelectual, iniciativa de aprendizagem continua.\n"
        )
    else:
        comp = (
            "P: Quais sao as suas principais competencias tecnicas relevantes para este cargo?\n"
            "A: Alinhamento com os requisitos, profundidade de conhecimento.\n"
            "P: Descreva uma situacao em que teve de aprender rapidamente uma nova ferramenta.\n"
            "A: Capacidade de aprendizagem, adaptacao, proactividade.\n"
            "P: Como garante a qualidade do seu trabalho tecnico?\n"
            "A: Metodo, atencao ao detalhe, orientacao para resultados.\n"
        )
    return f"""## APRESENTACAO E MOTIVACAO
P: Apresente-se brevemente e descreva o seu percurso profissional.
A: Capacidade de sintese, clareza de comunicacao, coerencia do percurso.
P: O que o/a motivou a candidatar-se a este cargo na nossa organizacao?
A: Conhecimento da organizacao, motivacao genuina, alinhamento de valores.

## EXPERIENCIA E HISTORIAL PROFISSIONAL
P: Descreva a sua experiencia mais relevante para o cargo de {vaga.titulo}.
A: Alinhamento com os requisitos da vaga, profundidade e qualidade da experiencia.
P: Qual foi o maior desafio profissional que enfrentou e como o resolveu?
A: Capacidade de resolucao de problemas, resiliencia, aprendizagem com a experiencia.

## COMPETENCIAS TECNICAS
{comp}
## COMPETENCIAS COMPORTAMENTAIS
P: Como gere situacoes de conflito com colegas ou superiores hierarquicos?
A: Inteligencia emocional, comunicacao assertiva, capacidade de mediar.
P: Descreva uma situacao em que teve de trabalhar sob pressao e com prazos apertados.
A: Resistencia ao stress, organizacao, capacidade de priorizar.

## SITUACOES HIPOTETICAS
P: Se tivesse de gerir varias tarefas urgentes em simultaneo, como procederia?
A: Gestao de prioridades, metodologia de trabalho, pedido de apoio quando necessario.

## QUESTOES DO CANDIDATO
P: Dar espaco ao/a candidato/a para colocar questoes sobre o cargo e a organizacao.
A: Interesse genuino, qualidade e pertinencia das questoes colocadas."""


def guiao_juri_view(request, token):
    """Public view for committee chair — no login required."""
    from .models import InterviewGuideSession
    session = get_object_or_404(InterviewGuideSession, token=token)

    if session.estado == InterviewGuideSession.ESTADO_APROVADO:
        return render(request, "vagas/guiao_juri_fechado.html", {"session": session})

    texto = session.texto_editado or session.texto_gerado
    categorias_parsed = _parse_perguntas(texto)
    categorias = [{"nome": nome, "perguntas": rows} for nome, rows in categorias_parsed]

    if request.method == "POST":
        cat_nomes = request.POST.getlist("cat_nome[]")
        perguntas = request.POST.getlist("pergunta[]")
        avaliar = request.POST.getlist("avaliar[]")
        cat_indices = request.POST.getlist("cat_index[]")

        linhas = []
        for nome in cat_nomes:
            linhas.append(f"## {nome}")
        texto_novo = _reconstruct_texto(cat_nomes, perguntas, avaliar, cat_indices)

        session.texto_editado = texto_novo
        session.estado = InterviewGuideSession.ESTADO_SUBMETIDO
        session.save()

        return render(request, "vagas/guiao_juri_obrigado.html", {"session": session})

    return render(request, "vagas/guiao_juri.html", {
        "session": session,
        "vaga": session.vaga,
        "categorias": categorias,
    })


def _reconstruct_texto(cat_nomes, perguntas, avaliar, cat_indices):
    """Rebuild P:/A: text from form POST data."""
    buckets = {i: [] for i in range(len(cat_nomes))}
    for i, (p, a) in enumerate(zip(perguntas, avaliar)):
        try:
            cat_i = int(cat_indices[i])
        except (IndexError, ValueError):
            cat_i = len(cat_nomes) - 1
        buckets[cat_i].append((p, a))

    lines = []
    for i, nome in enumerate(cat_nomes):
        lines.append(f"## {nome}")
        for p, a in buckets.get(i, []):
            if p.strip():
                lines.append(f"P: {p}")
                lines.append(f"A: {a}")
    return "\n".join(lines)


def guiao_aprovar(request, pk, session_id):
    """HR reviews and approves the chair's submitted guide."""
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    from .models import InterviewGuideSession
    session = get_object_or_404(InterviewGuideSession, pk=session_id, vaga=vaga)

    texto = session.texto_editado or session.texto_gerado
    categorias_parsed = _parse_perguntas(texto)
    categorias = [{"nome": nome, "perguntas": rows} for nome, rows in categorias_parsed]

    if request.method == "POST":
        action = request.POST.get("action")
        if action == "aprovar":
            # Save any HR edits before approving
            cat_nomes = request.POST.getlist("cat_nome[]")
            perguntas_list = request.POST.getlist("pergunta[]")
            avaliar_list = request.POST.getlist("avaliar[]")
            cat_indices = request.POST.getlist("cat_index[]")
            if cat_nomes and perguntas_list:
                session.texto_editado = _reconstruct_texto(cat_nomes, perguntas_list, avaliar_list, cat_indices)
            session.estado = InterviewGuideSession.ESTADO_APROVADO
            session.save()
            messages.success(request, "Guião aprovado. Pode agora fazer o download.")
            return redirect("guiao_download", pk=pk, session_id=session.pk)
        elif action == "devolver":
            session.estado = InterviewGuideSession.ESTADO_PENDENTE
            session.save()
            messages.info(request, "Guião devolvido ao presidente do júri para revisão.")
            return redirect("vaga_detail", pk=pk)

    return render(request, "vagas/guiao_aprovar.html", {
        "vaga": vaga,
        "session": session,
        "categorias": categorias,
    })


def guiao_download(request, pk, session_id):
    """Download approved guide as Word scoring table."""
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    from .models import InterviewGuideSession
    session = get_object_or_404(InterviewGuideSession, pk=session_id, vaga=vaga)

    if session.estado != InterviewGuideSession.ESTADO_APROVADO:
        messages.error(request, "O guião ainda não foi aprovado.")
        return redirect("vaga_detail", pk=pk)

    texto = session.texto_final()
    categorias = _parse_perguntas(texto)
    return _build_word_doc(vaga, categorias)


def _parse_perguntas(texto):
    """Parse structured interview guide text into list of (categoria, [(pergunta, avaliar)])."""
    categorias = []
    current_cat = None
    current_rows = []
    current_p = None

    for line in texto.split('\n'):
        line = line.strip()
        if line.startswith('## '):
            if current_cat is not None:
                if current_p:
                    current_rows.append((current_p, ''))
                    current_p = None
                categorias.append((current_cat, current_rows))
            current_cat = line[3:].strip()
            current_rows = []
            current_p = None
        elif line.startswith('P: ') or line.startswith('P:'):
            if current_p:
                current_rows.append((current_p, ''))
            current_p = line[2:].strip().lstrip(':').strip()
        elif line.startswith('A: ') or line.startswith('A:'):
            avaliar = line[2:].strip().lstrip(':').strip()
            if current_p:
                current_rows.append((current_p, avaliar))
                current_p = None
            else:
                current_rows.append(('', avaliar))

    if current_cat is not None:
        if current_p:
            current_rows.append((current_p, ''))
        categorias.append((current_cat, current_rows))

    return categorias


def _build_word_doc(vaga, categorias):
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    org_name = vaga.organisation.name if vaga.organisation else "Organizacao"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"GUIAO DE ENTREVISTA — {vaga.titulo.upper()}")
    run.bold = True
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = p.add_run(f"{org_name}   |   Data: ___/___/______   |   Candidato/a: _______________________________")
    sub.font.size = Pt(9)
    sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    COL_NUM   = Cm(0.8)
    COL_PERG  = Cm(8.5)
    COL_AVAL  = Cm(4.5)
    COL_SCORE = Cm(1.5)
    COL_NOTAS = Cm(2.2)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def cell_para(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color

    for cat_name, rows in categorias:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        for i, w in enumerate([COL_NUM, COL_PERG, COL_AVAL, COL_SCORE, COL_NOTAS]):
            table.columns[i].width = w

        hdr_row = table.rows[0]
        hdr_row.cells[0].merge(hdr_row.cells[4])
        hdr_cell = hdr_row.cells[0]
        set_cell_bg(hdr_cell, '1E3A5F')
        cell_para(hdr_cell, cat_name, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT, color=RGBColor(0xFF, 0xFF, 0xFF))

        col_row = table.add_row()
        for i, label in enumerate(['#', 'Pergunta', 'O que avaliar', '1–5', 'Notas']):
            set_cell_bg(col_row.cells[i], 'D0DCF0')
            cell_para(col_row.cells[i], label, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

        for idx, (pergunta, av) in enumerate(rows, 1):
            r = table.add_row()
            cell_para(r.cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_para(r.cells[1], pergunta)
            cell_para(r.cells[2], av, color=RGBColor(0x44, 0x55, 0x66))
            cell_para(r.cells[3], '', align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_para(r.cells[4], '')
            for cell in r.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(4)
                    para.paragraph_format.space_after = Pt(4)

        doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Pontuação: 1 = Muito fraco   2 = Fraco   3 = Adequado   4 = Bom   5 = Excelente")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_titulo = vaga.titulo.replace(' ', '_')
    filename = f"guiao_entrevista_{safe_titulo}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@require_POST
def comite_adicionar_avaliador(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    from .models import ComiteSession
    nome = request.POST.get("nome", "").strip()
    email = request.POST.get("email", "").strip()
    if not nome:
        messages.error(request, "O nome do avaliador é obrigatório.")
        return redirect("vaga_detail", pk=pk)
    if not email:
        messages.error(request, "O email do avaliador é obrigatório.")
        return redirect("vaga_detail", pk=pk)
    ComiteSession.objects.create(vaga=vaga, avaliador_nome=nome, avaliador_email=email)
    return redirect(f"/vagas/{pk}/?show_comite=1")


@require_POST
def comite_remover_avaliador(request, pk, session_pk):
    from .models import ComiteSession
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    session = get_object_or_404(ComiteSession, pk=session_pk, vaga=vaga)
    session.delete()
    messages.success(request, "Membro removido do comité.")
    return redirect("vaga_detail", pk=pk)


@require_POST
def comite_repor_avaliacao(request, pk, session_pk):
    from .models import ComiteSession, ComiteAvaliacao
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    session = get_object_or_404(ComiteSession, pk=session_pk, vaga=vaga)
    ComiteAvaliacao.objects.filter(session=session).delete()
    session.estado = ComiteSession.ESTADO_PENDENTE
    session.save(update_fields=["estado"])
    messages.success(request, f"Avaliação de {session.avaliador_nome} reposta. O membro pode voltar a preencher com o guião actualizado.")
    return redirect("vaga_detail", pk=pk)


def comite_avaliacao_view(request, token):
    from .models import ComiteSession, ComiteAvaliacao, InterviewGuideSession
    from candidatos.models import Candidato

    session = get_object_or_404(ComiteSession, token=token)
    vaga = session.vaga

    secoes = []
    perguntas = []
    guiao = (InterviewGuideSession.objects
             .filter(vaga=vaga, estado=InterviewGuideSession.ESTADO_APROVADO)
             .order_by("-created_at").first())
    if guiao:
        import re
        current_secao = {"titulo": "", "perguntas": []}
        for line in guiao.texto_final().splitlines():
            # Section heading: ## Heading or **Heading**
            heading = re.match(r'^#{1,3}\s*(.+)', line) or re.match(r'^\*{1,2}(.+?)\*{1,2}\s*$', line)
            if heading:
                if current_secao["perguntas"]:
                    secoes.append(current_secao)
                current_secao = {"titulo": heading.group(1).strip().rstrip(':'), "perguntas": []}
                continue
            # Format: "P: pergunta" (used by _reconstruct_texto)
            p_match = re.match(r'^P:\s*(.+)', line)
            if p_match:
                q = p_match.group(1).strip()
                if len(q) > 8:
                    current_secao["perguntas"].append(q)
                    perguntas.append(q)
                continue
            # Fallback: numbered questions "1. pergunta"
            q_match = re.match(r'^\s*\d+[\.\)]\s*(.+)', line)
            if q_match:
                q = q_match.group(1).strip()
                if len(q) > 8:
                    current_secao["perguntas"].append(q)
                    perguntas.append(q)
        if current_secao["perguntas"]:
            secoes.append(current_secao)
        if not secoes and perguntas:
            secoes = [{"titulo": "Perguntas de Entrevista", "perguntas": perguntas}]
        # Add flat index to each question so template can name fields c{pk}_resp_{idx}
        flat_idx = 0
        for s in secoes:
            qs_with_idx = []
            for q in s["perguntas"]:
                qs_with_idx.append({"texto": q, "idx": flat_idx})
                flat_idx += 1
            s["perguntas"] = qs_with_idx

    candidatos = Candidato.objects.filter(vaga=vaga, etapa="Entrevista").order_by("nome")

    if request.method == "POST":
        for c in candidatos:
            p = request.POST.get(f"c{c.pk}_pontuacao", "")
            ComiteAvaliacao.objects.update_or_create(
                session=session, candidato=c,
                defaults={
                    "pontuacao": int(p) if p.isdigit() and 1 <= int(p) <= 5 else None,
                    "recomendacao": request.POST.get(f"c{c.pk}_recomendacao", ""),
                    "pontos_fortes": request.POST.get(f"c{c.pk}_pontos_fortes", "").strip(),
                    "pontos_fracos": request.POST.get(f"c{c.pk}_pontos_fracos", "").strip(),
                    "notas": request.POST.get(f"c{c.pk}_notas", "").strip(),
                    "data_entrevista": request.POST.get(f"c{c.pk}_data_entrevista") or None,
                    "respostas_perguntas": [
                        {
                            "pergunta": q,
                            "nota": request.POST.get(f"c{c.pk}_resp_{i}", "").strip(),
                            "score": request.POST.get(f"c{c.pk}_qscore_{i}", ""),
                        }
                        for i, q in enumerate(perguntas)
                    ],
                }
            )
        session.estado = ComiteSession.ESTADO_SUBMETIDO
        session.save(update_fields=["estado"])
        return render(request, "vagas/comite_obrigado.html", {"session": session, "vaga": vaga})

    existing = {a.candidato_id: a for a in ComiteAvaliacao.objects.filter(session=session)}
    return render(request, "vagas/comite_avaliacao.html", {
        "session": session,
        "vaga": vaga,
        "candidatos": candidatos,
        "secoes": secoes,
        "perguntas": perguntas,
        "guiao": guiao,
        "existing": existing,
    })


@login_required
def comite_resultados(request, pk):
    from .models import ComiteSession, ComiteAvaliacao
    from candidatos.models import Candidato

    vaga = get_object_or_404(org_vagas(request), pk=pk)
    sessions = list(vaga.comite_sessions.prefetch_related("avaliacoes__candidato").order_by("created_at"))
    candidatos = list(Candidato.objects.filter(vaga=vaga, etapa="Entrevista").order_by("nome"))

    # Build matrix: candidato → list of (session, avaliacao|None)
    matrix = []
    for c in candidatos:
        row = []
        for s in sessions:
            av = next((a for a in s.avaliacoes.all() if a.candidato_id == c.pk), None)
            row.append({"session": s, "av": av})
        matrix.append({"candidato": c, "avaliacoes": row})

    return render(request, "vagas/comite_resultados.html", {
        "vaga": vaga,
        "sessions": sessions,
        "matrix": matrix,
    })


@require_POST
def comite_confirmar_decisoes(request, pk):
    from .models import ComiteSession, ComiteAvaliacao
    from candidatos.models import Candidato, NotaEntrevista
    from django.db import transaction
    from collections import defaultdict

    vaga = get_object_or_404(org_vagas(request), pk=pk)
    candidatos = Candidato.objects.filter(vaga=vaga, etapa="Entrevista")

    with transaction.atomic():
        for c in candidatos:
            avaliacoes = ComiteAvaliacao.objects.filter(
                session__vaga=vaga, session__estado=ComiteSession.ESTADO_SUBMETIDO, candidato=c
            )
            if not avaliacoes.exists():
                continue
            pontuacoes = [a.pontuacao for a in avaliacoes if a.pontuacao]
            media = round(sum(pontuacoes) / len(pontuacoes)) if pontuacoes else None
            recomendacoes = [a.recomendacao for a in avaliacoes if a.recomendacao]
            rec_final = max(set(recomendacoes), key=recomendacoes.count) if recomendacoes else ""
            notas_concat = "\n\n".join(
                f"[{a.session.avaliador_nome}]: {a.notas}" for a in avaliacoes if a.notas
            )
            NotaEntrevista.objects.update_or_create(
                candidato=c,
                defaults={"pontuacao": media, "recomendacao": rec_final, "notas": notas_concat}
            )
    messages.success(request, "Avaliações do comité consolidadas. Confirme agora a decisão final para cada candidato.")
    return redirect("vaga_detail", pk=pk)


@require_POST
def candidato_decisao_final(request, pk, candidato_pk):
    from candidatos.models import Candidato
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    candidato = get_object_or_404(Candidato, pk=candidato_pk, vaga=vaga)
    decisao = request.POST.get("decisao", "")
    etapa_map = {"contratar": "Proposta", "rejeitar": "Rejeitado", "considerar": "Entrevista"}
    nova_etapa = etapa_map.get(decisao)
    if nova_etapa:
        candidato.etapa = nova_etapa
        candidato.save(update_fields=["etapa"])
        messages.success(request, f"Decisão registada para {candidato.nome}.")
    return redirect("vaga_detail", pk=pk)


@require_POST
def enviar_avaliacao_grupo(request, pk):
    import uuid as _uuid
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    from candidatos.models import Candidato, AvaliacaoSession

    # Generate group token if not set
    if not vaga.avaliacao_group_token:
        vaga.avaliacao_group_token = _uuid.uuid4()
        vaga.save(update_fields=["avaliacao_group_token"])

    chair_email = request.POST.get("chair_email", "").strip()

    # Create AvaliacaoSession for every entrevista candidate that doesn't have one
    cands_entrevista = Candidato.objects.filter(
        vaga=vaga, organisation=request.user.organisation, etapa="Entrevista"
    )
    for c in cands_entrevista:
        if not AvaliacaoSession.objects.filter(candidato=c).exists():
            AvaliacaoSession.objects.create(candidato=c, chair_email=chair_email)
        elif chair_email:
            AvaliacaoSession.objects.filter(candidato=c).update(chair_email=chair_email)

    return redirect(f"/vagas/{pk}/?show_av_link=1")


def avaliacao_grupo_juri(request, token):
    from candidatos.models import AvaliacaoSession, Candidato
    from .models import InterviewGuideSession

    vaga = get_object_or_404(Vaga, avaliacao_group_token=token)

    # Get interview questions from approved guide
    perguntas = []
    guiao = (InterviewGuideSession.objects
             .filter(vaga=vaga, estado=InterviewGuideSession.ESTADO_APROVADO)
             .order_by("-created_at").first())
    if guiao:
        import re
        for m in re.finditer(r'^\s*\d+[\.\)]\s*(.+)', guiao.texto_final(), re.MULTILINE):
            q = m.group(1).strip()
            if len(q) > 10:
                perguntas.append(q)

    sessions = (AvaliacaoSession.objects
                .filter(candidato__vaga=vaga)
                .select_related("candidato")
                .order_by("candidato__nome"))

    if request.method == "POST":
        for session in sessions:
            if session.estado == AvaliacaoSession.ESTADO_CONFIRMADA:
                continue
            prefix = f"c{session.candidato.pk}_"
            p = request.POST.get(f"{prefix}pontuacao", "")
            session.pontuacao = int(p) if p.isdigit() and 1 <= int(p) <= 5 else None
            session.recomendacao = request.POST.get(f"{prefix}recomendacao", "")
            session.pontos_fortes = request.POST.get(f"{prefix}pontos_fortes", "").strip()
            session.pontos_fracos = request.POST.get(f"{prefix}pontos_fracos", "").strip()
            session.notas = request.POST.get(f"{prefix}notas", "").strip()
            session.data_entrevista = request.POST.get(f"{prefix}data_entrevista") or None
            respostas = []
            for i, q in enumerate(perguntas):
                nota = request.POST.get(f"{prefix}resp_{i}", "").strip()
                respostas.append({"pergunta": q, "nota": nota})
            session.respostas_perguntas = respostas
            session.estado = AvaliacaoSession.ESTADO_SUBMETIDA
            session.save()
        return render(request, "candidatos/avaliacao_obrigado.html", {"session": sessions.first()})

    return render(request, "vagas/avaliacao_grupo_juri.html", {
        "vaga": vaga,
        "sessions": sessions,
        "perguntas": perguntas,
    })


@require_POST
def confirmar_avaliacoes_grupo(request, pk):
    from candidatos.models import AvaliacaoSession, Candidato, NotaEntrevista
    from django.db import transaction

    vaga = get_object_or_404(org_vagas(request), pk=pk)
    sessions = AvaliacaoSession.objects.filter(
        candidato__vaga=vaga,
        estado=AvaliacaoSession.ESTADO_SUBMETIDA
    ).select_related("candidato")

    with transaction.atomic():
        for session in sessions:
            candidato = session.candidato
            AvaliacaoSession.objects.filter(pk=session.pk).update(
                estado=AvaliacaoSession.ESTADO_CONFIRMADA
            )
            NotaEntrevista.objects.update_or_create(
                candidato=candidato,
                defaults={
                    "data_entrevista": session.data_entrevista,
                    "pontuacao": session.pontuacao,
                    "recomendacao": session.recomendacao,
                    "pontos_fortes": session.pontos_fortes,
                    "pontos_fracos": session.pontos_fracos,
                    "notas": session.notas,
                }
            )
            rec = session.recomendacao
            nova_etapa = "Proposta" if rec == "recomendado" else ("Rejeitado" if rec == "nao_recomendado" else candidato.etapa)
            Candidato.objects.filter(pk=candidato.pk).update(etapa=nova_etapa)

    messages.success(request, "Avaliações confirmadas. Os candidatos foram actualizados.")
    return redirect("vaga_detail", pk=pk)


def shortlist(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    from candidatos.models import Candidato
    candidatos_triagem = (
        Candidato.objects
        .filter(vaga=vaga, etapa="Em Triagem")
        .order_by("-score_fit", "nome")
    )
    candidatos_entrevista = (
        Candidato.objects
        .filter(vaga=vaga, etapa="Entrevista")
        .order_by("-score_fit", "nome")
    )
    return render(request, "vagas/shortlist.html", {
        "vaga": vaga,
        "candidatos_triagem": candidatos_triagem,
        "candidatos_entrevista": candidatos_entrevista,
    })


@require_POST
def mover_para_entrevista(request, pk, candidato_pk):
    from candidatos.models import Candidato
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    candidato = get_object_or_404(Candidato, pk=candidato_pk, vaga=vaga, organisation=request.user.organisation)
    if candidato.etapa == "Em Triagem":
        Candidato.objects.filter(pk=candidato.pk).update(etapa="Entrevista")
        messages.success(request, f"{candidato.nome} movido para Entrevista.")
    return redirect("shortlist", pk=pk)


def relatorio_selecao(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    from candidatos.models import Candidato, NotaEntrevista

    candidatos_avaliados = (
        Candidato.objects
        .filter(vaga=vaga)
        .exclude(etapa__in=["Candidatura Recebida", "Em Triagem"])
        .select_related("nota_entrevista")
        .order_by("-nota_entrevista__pontuacao", "-score_fit", "nome")
    )

    # Build data for LLM + template
    dados_candidatos = []
    for c in candidatos_avaliados:
        nota = getattr(c, "nota_entrevista", None)
        dados_candidatos.append({
            "nome": c.nome,
            "etapa": c.etapa,
            "score_fit": c.score_fit,
            "pontuacao_entrevista": nota.pontuacao if nota else None,
            "recomendacao": nota.recomendacao if nota else "",
            "pontos_fortes": nota.pontos_fortes if nota else "",
            "pontos_fracos": nota.pontos_fracos if nota else "",
            "notas": nota.notas if nota else "",
            "experiencia_anos": c.experiencia_anos,
            "competencias": c.competencias,
        })

    narrativa = _gerar_narrativa_relatorio(vaga, dados_candidatos)

    return render(request, "vagas/relatorio_selecao.html", {
        "vaga": vaga,
        "candidatos": dados_candidatos,
        "narrativa": narrativa,
    })


def relatorio_selecao_download(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)
    from candidatos.models import Candidato
    import io
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    candidatos_avaliados = (
        Candidato.objects
        .filter(vaga=vaga)
        .exclude(etapa__in=["Candidatura Recebida", "Em Triagem"])
        .select_related("nota_entrevista")
        .order_by("-nota_entrevista__pontuacao", "-score_fit", "nome")
    )

    dados_candidatos = []
    for c in candidatos_avaliados:
        nota = getattr(c, "nota_entrevista", None)
        dados_candidatos.append({
            "nome": c.nome,
            "etapa": c.etapa,
            "score_fit": c.score_fit,
            "pontuacao_entrevista": nota.pontuacao if nota else None,
            "recomendacao": nota.recomendacao if nota else "",
            "pontos_fortes": nota.pontos_fortes if nota else "",
            "pontos_fracos": nota.pontos_fracos if nota else "",
            "notas": nota.notas if nota else "",
            "experiencia_anos": c.experiencia_anos,
            "competencias": c.competencias,
        })

    narrativa = _gerar_narrativa_relatorio(vaga, dados_candidatos)

    doc = Document()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Relatório de Seleção — {vaga.titulo}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    from datetime import date
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Elaborado em {date.today().strftime('%d/%m/%Y')}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # Narrativa IA
    if narrativa:
        h = doc.add_paragraph()
        rh = h.add_run("Análise e Recomendação")
        rh.bold = True
        rh.font.size = Pt(12)
        rh.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        for line in narrativa.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_paragraph()

    # Candidate summaries
    h2 = doc.add_paragraph()
    rh2 = h2.add_run("Resumo por Candidato")
    rh2.bold = True
    rh2.font.size = Pt(12)
    rh2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    REC_LABELS = {
        "recomendado": "Recomendado",
        "a_considerar": "A considerar",
        "nao_recomendado": "Não recomendado",
        "": "Sem avaliação",
    }
    STARS = {1: "★☆☆☆☆", 2: "★★☆☆☆", 3: "★★★☆☆", 4: "★★★★☆", 5: "★★★★★"}

    for idx, c in enumerate(dados_candidatos, 1):
        p_name = doc.add_paragraph()
        r_name = p_name.add_run(f"{idx}. {c['nome']}")
        r_name.bold = True
        r_name.font.size = Pt(11)

        meta_parts = []
        if c["pontuacao_entrevista"]:
            meta_parts.append(f"Entrevista: {STARS.get(c['pontuacao_entrevista'], '')} ({c['pontuacao_entrevista']}/5)")
        if c["score_fit"] is not None:
            meta_parts.append(f"Fit: {c['score_fit']}%")
        meta_parts.append(f"Etapa: {c['etapa']}")
        meta_parts.append(f"Recomendação: {REC_LABELS.get(c['recomendacao'], c['recomendacao'])}")
        doc.add_paragraph(" · ".join(meta_parts)).runs[0].font.size = Pt(9)

        if c["pontos_fortes"]:
            pf = doc.add_paragraph()
            pf.add_run("Pontos fortes: ").bold = True
            pf.add_run(c["pontos_fortes"])
        if c["pontos_fracos"]:
            pp = doc.add_paragraph()
            pp.add_run("Pontos fracos: ").bold = True
            pp.add_run(c["pontos_fracos"])
        if c["notas"]:
            pn = doc.add_paragraph()
            pn.add_run("Notas: ").bold = True
            pn.add_run(c["notas"])

        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    safe_titulo = vaga.titulo.replace(" ", "_")
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="relatorio_selecao_{safe_titulo}.docx"'
    return response


def _gerar_narrativa_relatorio(vaga, dados_candidatos):
    """Call LLM to generate a selection narrative. Returns text or None."""
    from core.llm import get_llm_response

    if not dados_candidatos:
        return None

    candidatos_text = ""
    REC_PT = {"recomendado": "Recomendado", "a_considerar": "A considerar", "nao_recomendado": "Não recomendado"}
    for c in dados_candidatos:
        rec = REC_PT.get(c["recomendacao"], "Sem avaliação de júri")
        pts = f"Pontuação entrevista: {c['pontuacao_entrevista']}/5. " if c["pontuacao_entrevista"] else ""
        fit = f"Fit com vaga: {c['score_fit']}%. " if c["score_fit"] is not None else ""
        fortes = f"Pontos fortes: {c['pontos_fortes']}. " if c["pontos_fortes"] else ""
        fracos = f"Pontos fracos: {c['pontos_fracos']}. " if c["pontos_fracos"] else ""
        candidatos_text += f"\n- {c['nome']} ({rec}): {pts}{fit}{fortes}{fracos}"

    req_comp = ", ".join(vaga.competencias_requeridas[:8]) if vaga.competencias_requeridas else "não especificadas"

    prompt = f"""Vaga: {vaga.titulo}
Requisitos: {vaga.nivel_formacao or ''}, {vaga.anos_experiencia_min or 0} anos experiência mínima
Competências requeridas: {req_comp}

Candidatos avaliados em entrevista:{candidatos_text}

Escreve um relatório de seleção profissional e conciso em Português europeu com:
1. Sumário executivo (2–3 frases sobre o processo)
2. Comparação dos candidatos (pontos diferenciadores)
3. Recomendação final clara (quem deve ser contratado e porquê)

Tom profissional, directo, sem repetir dados já listados. Máximo 300 palavras."""

    return get_llm_response(prompt, system="És um especialista em recrutamento e seleção de recursos humanos.")


def download_perguntas(request, pk):
    vaga = get_object_or_404(org_vagas(request), pk=pk)

    cat_nomes = request.POST.getlist("cat_nome[]")
    perguntas = request.POST.getlist("pergunta[]")
    avaliar_list = request.POST.getlist("avaliar[]")
    cat_indices = request.POST.getlist("cat_index[]")

    if cat_nomes and perguntas:
        texto = _reconstruct_texto(cat_nomes, perguntas, avaliar_list, cat_indices)
        categorias = _parse_perguntas(texto)
    else:
        texto = request.session.get(f"perguntas_{pk}", "")
        if not texto:
            return redirect("vaga_detail", pk=pk)
        categorias = _parse_perguntas(texto)

    if not categorias:
        return redirect("vaga_detail", pk=pk)

    return _build_word_doc(vaga, categorias)

