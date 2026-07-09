п»їimport os
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.views.decorators.http import require_POST
from django.contrib import messages
from vagas.models import Vaga
from candidatos.models import Candidato


def scoring_view(request):
    vagas = Vaga.objects.filter(organisation=request.user.organisation).order_by("-created_at")
    vaga_sel = None
    candidatos = []

    vaga_id = request.GET.get("vaga")
    if vaga_id:
        try:
            vaga_sel = Vaga.objects.get(pk=vaga_id, organisation=request.user.organisation)
            candidatos = Candidato.objects.filter(vaga=vaga_sel).order_by("-score_fit")
        except Vaga.DoesNotExist:
            pass

    sem_vaga_count = Candidato.objects.filter(
        organisation=request.user.organisation, vaga__isnull=True
    ).count()

    return render(request, "scoring/scoring.html", {
        "vagas": vagas,
        "vaga_sel": vaga_sel,
        "candidatos": candidatos,
        "sem_vaga_count": sem_vaga_count,
    })


def scoring_vaga(request, vaga_id):
    vaga = get_object_or_404(Vaga, pk=vaga_id, organisation=request.user.organisation)
    candidatos = Candidato.objects.filter(vaga=vaga).order_by("-score_fit")
    return redirect(f"/scoring/?vaga={vaga_id}")


@require_POST
def score_calculate(request):
    """HTMX Гўв‚¬вЂќ calculate fit score for one or all candidates of a vaga."""
    from django.conf import settings
    vaga_id = request.POST.get("vaga_id")
    candidato_id = request.POST.get("candidato_id")

    try:
        vaga = get_object_or_404(Vaga, pk=vaga_id, organisation=request.user.organisation)
    except Exception:
        return HttpResponse('<div class="alert-error">Vaga nГѓВЈo encontrada.</div>')

    os.environ["GROK_API_KEY"] = settings.GROK_API_KEY
    os.environ["LLM_ENGINE"] = settings.LLM_ENGINE

    from core.scorer import calcular_fit

    vaga_dict = {
        "titulo": vaga.titulo,
        "competencias_requeridas": vaga.competencias_requeridas or [],
        "anos_experiencia_min": vaga.anos_experiencia_min,
        "nivel_formacao": vaga.nivel_formacao,
        "responsabilidades": vaga.responsabilidades or [],
    }

    if candidato_id:
        qs = Candidato.objects.filter(pk=candidato_id, organisation=request.user.organisation)
    else:
        qs = Candidato.objects.filter(vaga=vaga, organisation=request.user.organisation)

    import logging
    log = logging.getLogger(__name__)
    from core.scorer import _score_deterministic

    for candidato in qs:
        cand_dict = {
            "nome": candidato.nome,
            "competencias": candidato.competencias or [],
            "experiencia_anos": candidato.experiencia_anos or 0,
            "formacao": candidato.formacao or [],
            "idiomas": candidato.idiomas or [],
            "resumo": candidato.resumo or "",
        }
        try:
            resultado = _score_deterministic(cand_dict, vaga_dict)
            resultado["metodo"] = "DeterminГ­stico"
            candidato.score_fit = resultado.get("score_total", 0)
            candidato.perfil_completo = resultado
            candidato.save(update_fields=["score_fit", "perfil_completo"])
        except Exception as e:
            log.error(f"Score error for {candidato.nome}: {e}", exc_info=True)

    return redirect(f"/scoring/?vaga={vaga_id}")


def exportar_excel(request):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    vaga_id = request.GET.get("vaga")
    vaga = get_object_or_404(Vaga, pk=vaga_id, organisation=request.user.organisation)
    candidatos = Candidato.objects.filter(vaga=vaga).order_by("-score_fit")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Ranking de Candidatos"

    # Styles
    header_fill = PatternFill("solid", fgColor="1D4ED8")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    title_font = Font(bold=True, size=14, color="1D4ED8")
    bold = Font(bold=True)
    center = Alignment(horizontal="center", vertical="center")
    thin = Side(style="thin", color="E2E8F0")
    border = Border(top=thin, bottom=thin, left=thin, right=thin)

    # Title block
    ws.merge_cells("A1:H1")
    ws["A1"] = f"Ranking de Candidatos Гўв‚¬вЂќ {vaga.titulo}"
    ws["A1"].font = title_font
    ws["A2"] = f"OrganizaГѓВ§ГѓВЈo: {vaga.organizacao or 'Гўв‚¬вЂќ'}   |   Local: {vaga.local or 'Гўв‚¬вЂќ'}   |   Candidatos: {candidatos.count()}"
    ws["A2"].font = Font(size=10, color="64748B")
    ws.append([])

    # Headers
    headers = ["#", "Nome", "Email", "ExperiГѓВЄncia (anos)", "Score Total", "CompetГѓВЄncias /50", "ExperiГѓВЄncia /30", "FormaГѓВ§ГѓВЈo /20", "Alinhamento", "Etapa Pipeline"]
    ws.append(headers)
    header_row = ws.max_row
    for col, _ in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = border

    # Data rows
    alt_fill = PatternFill("solid", fgColor="F8FAFC")
    green_fill = PatternFill("solid", fgColor="F0FDF4")
    for i, c in enumerate(candidatos, 1):
        det = c.perfil_completo or {}
        pts = det.get("pontuacao_detalhada") or {}
        row = [
            i,
            c.nome,
            c.email or "",
            c.experiencia_anos or 0,
            c.score_fit if c.score_fit is not None else "",
            pts.get("competencias", ""),
            pts.get("experiencia", ""),
            pts.get("formacao", ""),
            det.get("nivel_alinhamento", ""),
            c.etapa,
        ]
        ws.append(row)
        data_row = ws.max_row
        row_fill = green_fill if i == 1 else (alt_fill if i % 2 == 0 else None)
        for col in range(1, len(row) + 1):
            cell = ws.cell(row=data_row, column=col)
            cell.border = border
            cell.alignment = Alignment(vertical="center")
            if row_fill:
                cell.fill = row_fill
            if col == 5 and c.score_fit is not None:
                if c.score_fit >= 75:
                    cell.font = Font(bold=True, color="15803D")
                elif c.score_fit >= 50:
                    cell.font = Font(bold=True, color="92400E")
                else:
                    cell.font = Font(bold=True, color="DC2626")

    # Column widths
    widths = [5, 30, 30, 20, 14, 18, 18, 14, 18, 22]
    for col, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = w

    ws.row_dimensions[header_row].height = 20

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = f'attachment; filename="ranking_{vaga.titulo[:30]}.xlsx"'
    wb.save(response)
    return response


def exportar_word(request):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

    vaga_id = request.GET.get("vaga")
    vaga = get_object_or_404(Vaga, pk=vaga_id, organisation=request.user.organisation)
    candidatos = list(Candidato.objects.filter(vaga=vaga).order_by("-score_fit"))

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    def set_color(run, hex_color):
        run.font.color.rgb = RGBColor(
            int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        )

    def add_heading(text, level=1, color="1D4ED8"):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(
                int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            )
        return p

    def add_hr(doc):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CBD5E1")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(6)

    # ГўвЂќв‚¬ГўвЂќв‚¬ Cover / Title ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬ГўвЂќв‚¬
    title = doc.add_heading("", 0)
    title_run = title.add_run(вЂњRelatorio de Selecao de CandidatosвЂќ)
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    set_color(title_run, вЂњ1D4ED8вЂќ)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(fвЂќ{vaga.titulo}вЂќ)
    r.font.size = Pt(14)
    r.font.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = fвЂќData: {datetime.date.today().strftime('%d/%m/%Y')}   |   Candidatos avaliados: {len(candidatos)}вЂќ
    if vaga.organizacao:
        meta_text = fвЂќ{vaga.organizacao}   |   вЂњ + meta_text
    r2 = meta.add_run(meta_text)
    r2.font.size = Pt(10)
    set_color(r2, вЂњ64748BвЂќ)

    add_hr(doc)

    # 1. Sumario Executivo
    add_heading(вЂњ1. Sumario ExecutivoвЂќ, level=1)
    scored = [c for c in candidatos if c.score_fit is not None]
    alto = [c for c in scored if c.score_fit >= 75]
    medio = [c for c in scored if 50 <= c.score_fit < 75]
    baixo = [c for c in scored if c.score_fit < 50]

    if scored:
        top = scored[0]
        p = doc.add_paragraph()
        p.add_run(fвЂќForam avaliados {len(candidatos)} candidato(s) para a vaga de вЂњ).font.size = Pt(11)
        b = p.add_run(fвЂќ{vaga.titulo}вЂќ)
        b.bold = True; b.font.size = Pt(11)
        p.add_run(fвЂќ. {len(alto)} candidato(s) apresentam alinhamento Alto (>=75), {len(medio)} Medio (50-74) e {len(baixo)} Baixo (<50).вЂќ).font.size = Pt(11)

        p2 = doc.add_paragraph()
        p2.add_run(вЂњCandidato mais recomendado: вЂњ).font.size = Pt(11)
        b2 = p2.add_run(fвЂќ{top.nome} (Score: {top.score_fit}/100)вЂќ)
        b2.bold = True; b2.font.size = Pt(11)
        det_top = top.perfil_completo or {}
        if det_top.get(вЂњexplicacaoвЂќ):
            p2.add_run(fвЂќ - {det_top['explicacao'][0] if det_top['explicacao'] else ''}вЂќ).font.size = Pt(11)
    else:
        doc.add_paragraph(вЂњNenhum candidato foi avaliado com score de fit. Execute o calculo de scores antes de exportar o relatorio.вЂќ).font.size = Pt(11)

    add_hr(doc)

    # 2. Perfil da Vaga
    add_heading(вЂњ2. Perfil da VagaвЂќ, level=1)
    table_vaga = doc.add_table(rows=0, cols=2)
    table_vaga.style = вЂњTable GridвЂќ

    def add_vaga_row(label, value):
        if not value:
            return
        row = table_vaga.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = str(value)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[0].width = Cm(5)

    add_vaga_row(вЂњTituloвЂќ, vaga.titulo)
    add_vaga_row(вЂњOrganizacaoвЂќ, vaga.organizacao)
    add_vaga_row(вЂњLocalizacaoвЂќ, vaga.local)
    add_vaga_row(вЂњDepartamentoвЂќ, vaga.departamento)
    add_vaga_row(вЂњModalidadeвЂќ, vaga.modalidade)
    add_vaga_row(вЂњTipo de ContratoвЂќ, vaga.tipo_contrato)
    add_vaga_row(вЂњFormacao MinimaвЂќ, vaga.nivel_formacao)
    add_vaga_row(вЂњExperiencia MinimaвЂќ, fвЂќ{vaga.anos_experiencia_min} anosвЂќ if vaga.anos_experiencia_min else None)
    add_vaga_row(вЂњPrazoвЂќ, vaga.prazo_candidatura)
    if vaga.competencias_requeridas:
        add_vaga_row(вЂњCompetencias RequeridasвЂќ, вЂњ, вЂњ.join(vaga.competencias_requeridas))

    add_hr(doc)

    # 3. Ranking Geral
    add_heading(вЂњ3. Ranking Geral de CandidatosвЂќ, level=1)

    if candidatos:
        t = doc.add_table(rows=1, cols=6)
        t.style = вЂњTable GridвЂќ
        headers = [вЂњ#вЂќ, вЂњNomeвЂќ, вЂњScoreвЂќ, вЂњCompetenciasвЂќ, вЂњExperienciaвЂќ, вЂњFormacaoвЂќ]
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)
            set_color(run, вЂњFFFFFFвЂќ)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement(вЂњw:shdвЂќ)
            shd.set(qn(вЂњw:valвЂќ), вЂњclearвЂќ)
            shd.set(qn(вЂњw:colorвЂќ), вЂњautoвЂќ)
            shd.set(qn(вЂњw:fillвЂќ), вЂњ1D4ED8вЂќ)
            tcPr.append(shd)

        for i, c in enumerate(candidatos, 1):
            det = c.perfil_completo or {}
            pts = det.get(вЂњpontuacao_detalhadaвЂќ) or {}
            row = t.add_row()
            vals = [
                str(i),
                c.nome,
                fвЂќ{c.score_fit}/100вЂќ if c.score_fit is not None else вЂњ-вЂќ,
                fвЂќ{pts.get('competencias', '-')}/50вЂќ,
                fвЂќ{pts.get('experiencia', '-')}/30вЂќ,
                fвЂќ{pts.get('formacao', '-')}/20вЂќ,
            ]
            for j, v in enumerate(vals):
                row.cells[j].text = v
                run = row.cells[j].paragraphs[0].runs[0]
                run.font.size = Pt(10)
                if j == 2 and c.score_fit is not None:
                    run.bold = True
                    if c.score_fit >= 75:
                        set_color(run, вЂњ15803DвЂќ)
                    elif c.score_fit >= 50:
                        set_color(run, вЂњ92400EвЂќ)
                    else:
                        set_color(run, вЂњDC2626вЂќ)
    else:
        doc.add_paragraph(вЂњSem candidatos avaliados.вЂќ)

    add_hr(doc)

    # 4. Analise Individual
    add_heading(вЂњ4. Analise Individual por CandidatoвЂќ, level=1)

    for i, c in enumerate(candidatos, 1):
        det = c.perfil_completo or {}
        pts = det.get(вЂњpontuacao_detalhadaвЂќ) or {}

        p_name = doc.add_heading(вЂњвЂќ, level=2)
        r_rank = p_name.add_run(fвЂќ{i}. {c.nome}вЂќ)
        r_rank.font.size = Pt(13)
        set_color(r_rank, вЂњ1E293BвЂќ)

        score_p = doc.add_paragraph()
        score_p.paragraph_format.space_after = Pt(4)
        if c.score_fit is not None:
            r_score = score_p.add_run(fвЂќScore Fit: {c.score_fit}/100вЂќ)
            r_score.bold = True
            r_score.font.size = Pt(11)
            color = вЂњ15803DвЂќ if c.score_fit >= 75 else (вЂњ92400EвЂќ if c.score_fit >= 50 else вЂњDC2626вЂќ)
            set_color(r_score, color)
            nivel = det.get(вЂњnivel_alinhamentoвЂќ, вЂњвЂќ)
            if nivel:
                r_niv = score_p.add_run(fвЂќ   [{nivel}]вЂќ)
                r_niv.font.size = Pt(10)
                set_color(r_niv, color)
        else:
            score_p.add_run(вЂњScore Fit: Nao calculadoвЂќ).font.size = Pt(11)

        if pts:
            breakdown = doc.add_paragraph()
            breakdown.paragraph_format.space_after = Pt(2)
            b_run = breakdown.add_run(
                fвЂќCompetencias: {pts.get('competencias','-')}/50   |   вЂњ
                fвЂќExperiencia: {pts.get('experiencia','-')}/30   |   вЂњ
                fвЂќFormacao: {pts.get('formacao','-')}/20вЂќ
            )
            b_run.font.size = Pt(10)
            set_color(b_run, вЂњ475569вЂќ)

        info_parts = []
        if c.email:
            info_parts.append(fвЂќEmail: {c.email}вЂќ)
        if c.telefone:
            info_parts.append(fвЂќTel: {c.telefone}вЂќ)
        if c.experiencia_anos:
            info_parts.append(fвЂќExperiencia: {c.experiencia_anos} anosвЂќ)
        if c.etapa:
            info_parts.append(fвЂќPipeline: {c.etapa}вЂќ)
        if info_parts:
            info_p = doc.add_paragraph(вЂњ   |   вЂњ.join(info_parts))
            info_p.runs[0].font.size = Pt(10)
            set_color(info_p.runs[0], вЂњ64748BвЂќ)

        if c.competencias:
            skill_p = doc.add_paragraph()
            skill_r = skill_p.add_run(вЂњCompetencias: вЂњ)
            skill_r.bold = True
            skill_r.font.size = Pt(10)
            skill_p.add_run(вЂњ, вЂњ.join(c.competencias)).font.size = Pt(10)

        if c.formacao:
            form_p = doc.add_paragraph()
            form_r = form_p.add_run(вЂњFormacao: вЂњ)
            form_r.bold = True
            form_r.font.size = Pt(10)
            form_p.add_run(вЂњ / вЂњ.join(c.formacao)).font.size = Pt(10)

        if c.resumo:
            sum_p = doc.add_paragraph()
            sum_r = sum_p.add_run(вЂњResumo Profissional: вЂњ)
            sum_r.bold = True
            sum_r.font.size = Pt(10)
            sum_p.add_run(c.resumo[:600]).font.size = Pt(10)

        explicacao = det.get(вЂњexplicacaoвЂќ) or []
        if explicacao:
            expl_head = doc.add_paragraph()
            expl_r = expl_head.add_run(вЂњAnalise de Alinhamento:вЂќ)
            expl_r.bold = True
            expl_r.font.size = Pt(10)
            set_color(expl_r, вЂњ1D4ED8вЂќ)
            for note in explicacao:
                p_note = doc.add_paragraph(style=вЂќList BulletвЂќ)
                p_note.add_run(note).font.size = Pt(10)
                p_note.paragraph_format.space_after = Pt(1)

        rec_p = doc.add_paragraph()
        rec_r = rec_p.add_run(вЂњRecomendacao: вЂњ)
        rec_r.bold = True
        rec_r.font.size = Pt(10)
        if c.score_fit is not None:
            if c.score_fit >= 75:
                rec_text = вЂњCandidato altamente recomendado para avanГ§ar no processo de selecao. O perfil demonstra forte alinhamento com os requisitos da vaga.вЂќ
            elif c.score_fit >= 50:
                rec_text = вЂњCandidato com alinhamento moderado. Recomenda-se entrevista para verificar competencias especificas nao confirmadas pelo CV.вЂќ
            else:
                rec_text = вЂњCandidato com baixo alinhamento com os requisitos da vaga. Considerar apenas se houver escassez de outros candidatos.вЂќ
        else:
            rec_text = вЂњScore nao calculado. Execute a analise para obter uma recomendacao fundamentada.вЂќ
        rec_p.add_run(rec_text).font.size = Pt(10)

        if i < len(candidatos):
            doc.add_paragraph()
            add_hr(doc)

    add_hr(doc)

    # 5. Conclusao e Proximos Passos
    add_heading(вЂњ5. Conclusao e Proximos PassosвЂќ, level=1)
    conc = doc.add_paragraph()
    conc.add_run(
        fвЂќCom base na analise de {len(candidatos)} candidato(s) para a vaga de {vaga.titulo}, вЂњ
    ).font.size = Pt(11)
    if alto:
        b3 = conc.add_run(fвЂќrecomenda-se avanГ§ar com {', '.join([c.nome for c in alto[:3]])} вЂњ)
        b3.bold = True; b3.font.size = Pt(11)
        conc.add_run(вЂњpara a fase de entrevistas.вЂќ).font.size = Pt(11)
    else:
        conc.add_run(вЂњnenhum candidato atingiu o limiar de alta recomendacao (75+). Considere alargar o processo de recrutamento.вЂќ).font.size = Pt(11)

    steps = doc.add_paragraph()
    steps.add_run(вЂњPassos sugeridos:вЂќ).bold = True
    steps.runs[0].font.size = Pt(11)
    for step in [
        вЂњAgendar entrevistas com candidatos de score Alto (>=75)вЂќ,
        вЂњRever candidatos de score Medio (50-74) apos entrevistas iniciaisвЂќ,
        вЂњActualizar o pipeline no TalentIQ apos cada fase de selecaoвЂќ,
        вЂњDocumentar feedback das entrevistas nas notas de cada candidatoвЂќ,
    ]:
        p_s = doc.add_paragraph(style=вЂќList BulletвЂќ)
        p_s.add_run(step).font.size = Pt(10)
        p_s.paragraph_format.space_after = Pt(2)

    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(24)
    fr = footer_p.add_run(fвЂќRelatorio gerado pelo TalentIQ - {datetime.date.today().strftime('%d/%m/%Y')}вЂќ)
    fr.font.size = Pt(9)
    set_color(fr, вЂњ94A3B8вЂќ)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"relatorio_{vaga.titulo[:30].replace(' ', '_')}.docx"
    response = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def scoring_geral(request):
    candidatos = Candidato.objects.filter(
        organisation=request.user.organisation,
    ).select_related("vaga").order_by("-score_fit")
    return render(request, "scoring/geral.html", {"candidatos": candidatos})

